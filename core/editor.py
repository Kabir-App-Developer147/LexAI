import os
import shutil
from docx import Document

class FileEditor:
    def __init__(self, data_folder="data"):
        self.data_folder = data_folder

    def _get_copy_path(self, filename):
        name, ext = os.path.splitext(filename)
        return os.path.join(self.data_folder, f"{name}_modified{ext}")

    def edit_text_file(self, filename, new_content):
        """Creates a copy and replaces content for .txt or .md files."""
        src = os.path.join(self.data_folder, filename)
        dest = self._get_copy_path(filename)
        
        if not os.path.exists(src):
            return f"Error: File {filename} not found."
        
        try:
            with open(dest, 'w', encoding='utf-8') as f:
                f.write(new_content)
            return f"Success: Created modified copy at {os.path.basename(dest)}"
        except Exception as e:
            return f"Error editing file: {e}"

    def edit_docx_file(self, filename, new_content):
        """
        Creates a copy and appends/replaces content for .docx files.
        For now, this simple version appends the new content as a new paragraph.
        """
        src = os.path.join(self.data_folder, filename)
        dest = self._get_copy_path(filename)
        
        if not os.path.exists(src):
            return f"Error: File {filename} not found."
            
        try:
            doc = Document(src)
            # Simple implementation: append new content
            doc.add_paragraph("\n--- LexAI MODIFICATION ---\n")
            doc.add_paragraph(new_content)
            doc.save(dest)
            return f"Success: Created modified copy at {os.path.basename(dest)}"
        except Exception as e:
            return f"Error editing docx: {e}"

    def apply_to_main(self, filename):
        """Replaces the main file with the modified copy."""
        copy_name = f"{os.path.splitext(filename)[0]}_modified{os.path.splitext(filename)[1]}"
        src = os.path.join(self.data_folder, copy_name)
        dest = os.path.join(self.data_folder, filename)
        
        if not os.path.exists(src):
            return f"Error: No modified copy found for {filename}"
            
        try:
            shutil.move(src, dest)
            return f"Success: {filename} has been updated with your changes."
        except Exception as e:
            return f"Error applying changes: {e}"

_editor_instance = None

def get_editor():
    global _editor_instance
    if _editor_instance is None:
        _editor_instance = FileEditor()
    return _editor_instance
